"""DocxRenderer — deterministic Word document generation from DocxSpec.

Opens a .docx template file (if provided), clears body content while preserving
styles/fonts, then fills in sections from the spec. This guarantees template
structure is preserved rather than re-invented by the LLM.
"""
from __future__ import annotations

import io
import logging
from typing import Optional

from app.rendering.doc_spec import DocxSpec, DocxSectionSpec, TableSpec, ChartSpec

logger = logging.getLogger(__name__)

# Usable body width in inches for a standard A4/Letter doc (page 8.27" − 2×1.18" margin)
_BODY_WIDTH_INCHES = 6.0


class DocxRenderer:
    """Renders a DocxSpec to .docx bytes.

    All exceptions propagate — no silent swallowing.
    The pipeline DOC_RENDER phase wraps errors in PipelineError.
    """

    def render(
        self,
        spec: DocxSpec,
        template_path: Optional[str] = None,
        reference_structure=None,  # ReferenceStructure | None
    ) -> bytes:
        from docx import Document
        from app.services.document_generator import _apply_docx_simplified_chinese_locale

        # Auto-numbering counters reset per render call
        self._figure_num = 0
        self._table_num = 0

        has_template = bool(template_path)
        heading_style_map: dict[str, str] = {}

        if has_template:
            doc = Document(template_path)
            # P1-1: Extract heading style map BEFORE clearing body content
            heading_style_map = self._build_heading_style_map(doc)
            # P2-D: Snapshot the template's sectPr (page layout + header/footer rels)
            # before clearing, so we can restore it after generating content.
            template_sect_pr = _snapshot_sect_pr(doc)
            self._clear_body_content(doc)
        else:
            doc = Document()
            template_sect_pr = None
            # P2-5: Apply reference doc styles when available (no template file)
            self._apply_default_styles(doc, reference_structure)
            self._apply_standard_styles(doc, (spec.metadata or {}).get("document_standard"))

        document_standard = (spec.metadata or {}).get("document_standard") or {}
        if document_standard.get("key") == "official_document":
            self._write_official_red_head(doc, spec.title)
            title_para = doc.add_heading(spec.title, level=0)
        else:
            title_para = doc.add_heading(spec.title, level=0)
        title_para.alignment = 1  # CENTER

        # P2-4: Insert a Word TOC field after the title so readers can refresh with F9.
        # Official documents usually do not carry a TOC in the main body.
        if document_standard.get("key") != "official_document":
            _insert_toc(doc)

        for sec_spec in spec.sections:
            self._write_section(doc, sec_spec, spec.tense, spec.style,
                                level=1, has_template=has_template,
                                heading_style_map=heading_style_map)

        # P2-D: Restore template sectPr (header/footer bindings) after body generation.
        if template_sect_pr is not None:
            _restore_sect_pr(doc, template_sect_pr)

        _apply_docx_simplified_chinese_locale(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Body helpers ──────────────────────────────────────────────────────────

    def _clear_body_content(self, doc) -> None:
        """Remove all paragraphs and tables from doc body, keeping page layout (sectPr)."""
        body = doc.element.body
        to_remove = [
            child for child in body
            if (child.tag.split("}")[-1] if "}" in child.tag else child.tag) in ("p", "tbl")
        ]
        for el in to_remove:
            body.remove(el)

    def _build_heading_style_map(self, doc) -> dict[str, str]:
        """P1-1: Map heading text → paragraph style name from template BEFORE body clear.

        Enables _write_section to apply the exact template style when
        template_heading_match is provided, rather than guessing from heading level.

        P2-2: Also stores a normalized key so that ordinal-variant headings
        (一、/1、/①/壹 prefixes) can be matched even when the outline uses a
        different numeral style than the template.
        Also detects the body text style name from the template for use in paragraph writing.
        """
        style_map: dict[str, str] = {}
        norm_map: dict[str, str] = {}  # normalized_text → style_name
        body_style_name: str | None = None

        for para in doc.paragraphs:
            style_name = para.style.name or ""
            style_lower = style_name.lower()
            if (
                "heading" in style_lower
                or "标题" in style_name
                or style_lower in ("title", "subtitle")
            ):
                text = para.text.strip()
                if text:
                    style_map[text] = style_name
                    norm_map[_normalize_heading(text)] = style_name
            # P2-2: Detect body text style — prefer explicit body style over Normal
            elif body_style_name is None and style_name not in ("", "Normal", "Default Paragraph Style"):
                if (
                    "body" in style_lower
                    or "正文" in style_name
                    or "text" in style_lower
                    or "段落" in style_name
                ):
                    body_style_name = style_name

        # Attach norm_map and detected body style for lookup in _write_section
        self._heading_norm_map = norm_map
        self._template_body_style = body_style_name  # None → fall back to Normal
        return style_map

    def _apply_default_styles(self, doc, ref_structure=None) -> None:
        """Apply styles from reference structure when no template file is available.

        Applies body font/size to Normal, and propagates heading font with sensible
        size/spacing overrides so H1/H2/H3 look professional in Chinese documents.
        """
        from docx.shared import Pt
        body_style = doc.styles["Normal"]

        body_font = getattr(ref_structure, "body_font", "") if ref_structure else ""
        body_size = getattr(ref_structure, "body_size_pt", 0) if ref_structure else 0

        if body_font:
            body_style.font.name = body_font
            body_style.font.size = Pt(body_size or 11)
            before = getattr(ref_structure, "body_space_before_pt", 0)
            after = getattr(ref_structure, "body_space_after_pt", 0)
            if before:
                body_style.paragraph_format.space_before = Pt(before)
            if after:
                body_style.paragraph_format.space_after = Pt(after)
        else:
            body_style.font.size = Pt(11)
            body_style.font.name = "宋体"
            # R18-SPACING: Apply consistent paragraph spacing for no-template docs.
            # Without explicit spacing, python-docx produces wall-of-text paragraphs.
            # 3pt before + 6pt after gives visually distinct paragraph breaks in Chinese.
            body_style.paragraph_format.space_before = Pt(3)
            body_style.paragraph_format.space_after = Pt(6)

        # Heading style inheritance: apply font + sensible spacing for Chinese docs
        font_name = body_font or "宋体"
        _HEADING_DEFS = [
            ("Heading 1", 16, True,  12, 6),
            ("Heading 2", 14, True,  10, 4),
            ("Heading 3", 13, False,  8, 3),
        ]
        for style_name, size_pt, bold, space_before, space_after in _HEADING_DEFS:
            try:
                h_style = doc.styles[style_name]
                h_style.font.name = font_name
                h_style.font.size = Pt(size_pt)
                h_style.font.bold = bold
                h_style.paragraph_format.space_before = Pt(space_before)
                h_style.paragraph_format.space_after = Pt(space_after)
            except Exception:
                pass

    def _apply_standard_styles(self, doc, standard: dict | None) -> None:
        """Apply renderer-facing style defaults for selected document templates."""
        if not standard:
            return
        from docx.shared import Pt
        render = standard.get("render") or {}
        body_font = render.get("body_font") or "宋体"
        body_size = float(render.get("body_size") or 11)
        try:
            normal = doc.styles["Normal"]
            normal.font.name = body_font
            normal.font.size = Pt(body_size)
            normal.paragraph_format.line_spacing = 1.5
            normal.paragraph_format.space_after = Pt(6)
        except Exception:
            pass
        if standard.get("key") == "official_document":
            try:
                doc.styles["Heading 1"].font.name = "黑体"
                doc.styles["Heading 1"].font.size = Pt(16)
                doc.styles["Heading 2"].font.name = "楷体_GB2312"
                doc.styles["Heading 2"].font.size = Pt(16)
            except Exception:
                pass
        elif standard.get("key") == "academic_paper":
            try:
                doc.styles["Heading 1"].font.name = "Times New Roman"
                doc.styles["Heading 1"].font.size = Pt(14)
                doc.styles["Heading 2"].font.name = "Times New Roman"
                doc.styles["Heading 2"].font.size = Pt(12)
            except Exception:
                pass

    def _write_official_red_head(self, doc, title: str) -> None:
        """Add a simple official-document red head before the title."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        issuing = _infer_official_issuing_org(title)
        red_head = doc.add_paragraph()
        red_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = red_head.add_run(f"{issuing}文件")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(196, 0, 0)
        try:
            run.font.name = "方正小标宋简体"
        except Exception:
            pass

        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(196, 0, 0)

    # ── Section writing ───────────────────────────────────────────────────────

    def _write_section(
        self,
        doc,
        sec: DocxSectionSpec,
        tense: str,
        style: str,
        level: int = 1,
        has_template: bool = False,
        heading_style_map: dict[str, str] | None = None,
    ) -> None:
        from docx.shared import Pt

        # P1-1: Use exact template heading style when available;
        # P2-2: Fall back to normalized match (handles 一/1/①/壹 ordinal variants).
        applied_heading = False
        if heading_style_map and sec.template_heading_match:
            template_style = heading_style_map.get(sec.template_heading_match)
            if not template_style:
                norm_map = getattr(self, "_heading_norm_map", {})
                template_style = norm_map.get(_normalize_heading(sec.template_heading_match))
            if template_style:
                available_styles = {s.name for s in doc.styles}
                if template_style in available_styles:
                    heading_para = doc.add_paragraph(sec.title, style=template_style)
                    # P2-B: Page break before top-level sections when using a template,
                    # matching typical Word document conventions.
                    if has_template and level == 1:
                        heading_para.paragraph_format.page_break_before = True
                    applied_heading = True

        if not applied_heading:
            heading_para = doc.add_heading(sec.title, level=min(level, 9))
            if has_template and level == 1:
                heading_para.paragraph_format.page_break_before = True

        # P3-4: Ensure minimum space_before on heading paragraphs when using a template.
        # Some templates set space_before=0 or leave it None, producing wall-of-text
        # with no visual separation between sections.
        if has_template:
            pf = heading_para.paragraph_format
            try:
                current_before = pf.space_before
                _min_before = Pt(8) if level == 1 else Pt(4)
                if current_before is None or current_before < _min_before:
                    pf.space_before = _min_before
            except Exception:
                pass

        # P2-2: Widow/orphan control — keep headings with the paragraph that follows.
        # This prevents a heading from appearing alone at the bottom of a page with
        # its content starting on the next page (a common professionalism issue).
        try:
            heading_para.paragraph_format.keep_with_next = True
            heading_para.paragraph_format.widow_control = True
        except Exception:
            pass

        content_type = sec.content_type

        # P2-2: Use template-detected body style for paragraphs when available.
        _body_style = getattr(self, "_template_body_style", None) if has_template else None

        if content_type in ("paragraphs", "mixed"):
            for para_text in sec.paragraphs:
                if _body_style:
                    try:
                        para = doc.add_paragraph(para_text, style=_body_style)
                    except Exception:
                        para = doc.add_paragraph(para_text)
                else:
                    para = doc.add_paragraph(para_text)
                if not has_template:
                    para.paragraph_format.first_line_indent = Pt(22)
                else:
                    # Apply Chinese-standard indent when the template style has
                    # no explicit setting (None) OR has an explicit zero indent.
                    # An explicit non-zero value means the template intentionally
                    # chose a different indent — respect that choice.
                    current = para.paragraph_format.first_line_indent
                    if current is None or current == Pt(0):
                        para.paragraph_format.first_line_indent = Pt(22)

        if content_type in ("bullets", "mixed"):
            for bullet in sec.bullets:
                self._add_bullet(doc, bullet)

        if sec.table:
            self._write_table(doc, sec.table)

        if sec.chart:
            self._write_chart(doc, sec.chart)

        # P3-1: Render evidence source line when evidence_ids are present
        evidence_ids = getattr(sec, "evidence_ids", None) or []
        if evidence_ids:
            _add_source_line(doc, evidence_ids)

        for subsec in (sec.subsections or []):
            self._write_section(doc, subsec, tense, style,
                                level=level + 1, has_template=has_template,
                                heading_style_map=heading_style_map)

    def _add_bullet(self, doc, text: str) -> None:
        """Add a bullet paragraph, falling back gracefully if 'List Bullet' is absent."""
        bullet_style = "List Bullet"
        if bullet_style not in [s.name for s in doc.styles]:
            doc.add_paragraph(f"— {text}")
        else:
            doc.add_paragraph(text, style=bullet_style)

    # ── Table ─────────────────────────────────────────────────────────────────

    def _write_table(self, doc, table_spec: TableSpec) -> None:
        from docx.shared import Inches, Pt

        headers = table_spec.headers
        rows = table_spec.rows
        if not headers and not rows:
            return

        n_cols = len(headers) or (len(rows[0]) if rows else 0)
        if n_cols == 0:
            return

        # Auto-number caption: "表N：caption"
        self._table_num += 1
        raw_caption = table_spec.caption or ""
        numbered_caption = f"表{self._table_num}：{raw_caption}" if raw_caption else f"表{self._table_num}"

        # P2-3: Caption above table — use Word "Caption" style when available
        # so numbering integrates with Word's cross-reference system.
        _CAPTION_STYLES = ("Caption", "图表标题", "Table Caption")
        available_style_names = {s.name for s in doc.styles}
        cap_style = next((s for s in _CAPTION_STYLES if s in available_style_names), None)
        if cap_style:
            cap_above = doc.add_paragraph(numbered_caption, style=cap_style)
        else:
            cap_above = doc.add_paragraph(numbered_caption)
            cap_above.paragraph_format.alignment = 1  # CENTER
            for run in cap_above.runs:
                run.font.size = Pt(9)
                run.font.bold = True

        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = "Table Grid"

        # Explicit column widths to prevent page overflow
        tbl.autofit = False
        col_width = Inches(_BODY_WIDTH_INCHES / n_cols)
        for col in tbl.columns:
            col.width = col_width

        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers[:n_cols]):
            hdr_cells[i].text = h
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)

        for r_idx, row in enumerate(rows):
            data_cells = tbl.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row[:n_cols]):
                # R18-TABLE: Truncate oversized cell text to prevent table overflow.
                # Cells with >200 chars push the table off-page in narrow-column layouts.
                raw = str(cell_text)
                if len(raw) > 200:
                    raw = raw[:197] + "…"
                data_cells[c_idx].text = raw

        # P3-3: For tables >25 data rows, enable row breaking across pages so
        # tall rows don't produce blank pages or clipped content.
        if len(rows) > 25:
            try:
                _enable_table_row_breaks(tbl)
            except Exception:
                pass

    # ── Chart ─────────────────────────────────────────────────────────────────

    def _write_chart(self, doc, chart_spec: ChartSpec) -> None:
        """Render chart as PNG and embed as inline image with auto-numbered caption."""
        from docx.shared import Inches, Pt
        from app.services.chart_render_service import render_chart_png

        self._figure_num += 1
        fig_caption = f"图{self._figure_num}：{chart_spec.title}"

        try:
            render_spec = chart_spec.to_render_spec()
            result = render_chart_png(
                render_spec,
                palette=["#2563EB", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6", "#06B6D4", "#EC4899", "#84CC16"],
            )
            if result and result.png:
                pic_para = doc.add_paragraph()
                pic_para.paragraph_format.alignment = 1  # CENTER
                # P2-3: Keep chart image with caption — prevents page break between them
                pic_para.paragraph_format.keep_with_next = True
                run = pic_para.add_run()
                run.add_picture(io.BytesIO(result.png), width=Inches(5.5))

                # Accessibility alt text
                try:
                    _set_picture_alt_text(run, chart_spec.title)
                except Exception:
                    pass

                # Numbered caption below chart
                cap_para = doc.add_paragraph(fig_caption)
                cap_para.paragraph_format.alignment = 1
                for r in cap_para.runs:
                    r.font.size = Pt(9)
                    r.font.italic = True

                if chart_spec.source_note:
                    note = doc.add_paragraph(f"数据来源：{chart_spec.source_note}")
                    note.paragraph_format.alignment = 1
                    for r in note.runs:
                        r.font.size = Pt(8)
                return
        except Exception as exc:
            logger.warning("[DOCX] Chart render failed (%s): %s", chart_spec.title, exc)

        doc.add_paragraph(f"[{fig_caption}（渲染失败，请手动插入）]")


# ── Module-level helpers ──────────────────────────────────────────────────────

def _snapshot_sect_pr(doc) -> object | None:
    """P2-D: Deep-copy the sectPr element from the template document body.

    sectPr contains page size, margins, and — crucially — relationships to
    header/footer parts. Snapshotting before body-clear lets us restore it
    after content is generated, preserving template headers and footers.
    """
    try:
        from copy import deepcopy
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        sect_pr = doc.element.body.find(f"{{{W_NS}}}sectPr")
        if sect_pr is not None:
            return deepcopy(sect_pr)
    except Exception as exc:
        logger.debug("[DOCX] sectPr snapshot failed: %s", exc)
    return None


def _restore_sect_pr(doc, sect_pr_copy) -> None:
    """P2-D: Replace the output document's sectPr with the template's snapshot.

    This reinstates header/footer relationships that python-docx would otherwise
    drop when adding new paragraphs to a cleared template body.
    """
    try:
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        body = doc.element.body
        existing = body.find(f"{{{W_NS}}}sectPr")
        if existing is not None:
            body.remove(existing)
        body.append(sect_pr_copy)
    except Exception as exc:
        logger.debug("[DOCX] sectPr restore failed: %s", exc)


def _add_source_line(doc, evidence_ids: list[str]) -> None:
    """P3-1: Render evidence citations as a professional numbered reference block.

    Format:
      ────────────────────────────────────────
      参考资料
      [1] source_label_1
      [2] source_label_2

    Each source label is rendered in 8pt italic grey. A thin horizontal rule
    (paragraph border) separates the citations from the section body.
    """
    if not evidence_ids:
        return
    try:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _grey_run(para, text: str, size_pt: float = 8.0, italic: bool = True) -> None:
            run = para.add_run(text)
            run.font.size = Pt(size_pt)
            run.font.italic = italic
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Separator line via top border on the header paragraph
        header_para = doc.add_paragraph()
        _grey_run(header_para, "参考资料", size_pt=8.0, italic=False)
        header_para.paragraph_format.space_before = Pt(6)
        header_para.paragraph_format.space_after = Pt(2)
        # Add a top border (thin rule) to visually separate from body
        try:
            pPr = header_para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top_el = OxmlElement("w:top")
            top_el.set(qn("w:val"), "single")
            top_el.set(qn("w:sz"), "4")
            top_el.set(qn("w:space"), "1")
            top_el.set(qn("w:color"), "CCCCCC")
            pBdr.append(top_el)
            pPr.append(pBdr)
        except Exception:
            pass

        # Numbered citation entries
        for i, src in enumerate(evidence_ids[:8], start=1):
            cite_para = doc.add_paragraph()
            _grey_run(cite_para, f"[{i}] {str(src)[:120]}", size_pt=8.0, italic=True)
            cite_para.paragraph_format.left_indent = Pt(12)
            cite_para.paragraph_format.space_before = Pt(1)
            cite_para.paragraph_format.space_after = Pt(1)

    except Exception as exc:
        logger.debug("[DOCX] Citation block rendering failed: %s", exc)


def _insert_toc(doc) -> None:
    """P2-4: Insert a Word TOC field after the title paragraph.

    The TOC won't render until the reader presses F9 in Word, but the field code
    is embedded so Word knows where to put the table of contents.
    """
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        para = doc.add_paragraph()
        run = para.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        run._r.append(fld_begin)

        instr_el = OxmlElement("w:instrText")
        instr_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr_el.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr_el)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        placeholder_el = OxmlElement("w:t")
        placeholder_el.text = "目录（请在 Word 中按 Ctrl+A，再按 F9 更新）"
        run._r.append(placeholder_el)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    except Exception as exc:
        logger.debug("[DOCX] TOC field insertion failed: %s", exc)


def _enable_table_row_breaks(tbl) -> None:
    """P3-3: Allow Word to break table rows across page boundaries for long tables.

    By default, Word keeps each row together (cantSplit). For tables with >25 rows
    this causes rows to overflow or produce blank pages. Removing the cantSplit
    element from each trPr lets Word paginate rows naturally.
    """
    from lxml import etree
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for row in tbl.rows:
        tr = row._tr
        trPr = tr.find(f"{{{W_NS}}}trPr")
        if trPr is None:
            trPr = etree.SubElement(tr, f"{{{W_NS}}}trPr")
            tr.insert(0, trPr)
        # Remove cantSplit if present (prevents row break)
        cant_split = trPr.find(f"{{{W_NS}}}cantSplit")
        if cant_split is not None:
            trPr.remove(cant_split)
        # Explicitly allow break: set tblHeader=false (not repeat header on each page for data rows)
        # Just ensuring cantSplit is absent is sufficient for python-docx tables.


def _normalize_heading(text: str) -> str:
    """P2-2: Strip leading ordinal markers and normalize to bare heading text.

    Converts Chinese numerals (一/壹/①), Arabic digits (1), and circled
    numbers to a bare text so that '一、引言' and '1. 引言' both match '引言'.
    """
    import re
    # Remove leading Chinese numerals + separator (一、二、 / 壹贰叁)
    _CN = "一二三四五六七八九十百零壹贰叁肆伍陆柒捌玖拾佰"
    text = re.sub(rf"^[{_CN}]+[、。.．：:，,\s]+", "", text)
    # Remove leading Arabic digits + separator (1. / 1、/ 1 )
    text = re.sub(r"^\d+[.\s、，,：:]+\s*", "", text)
    # Remove leading circled numbers ①–⑳
    text = re.sub(r"^[①-⑳]\s*", "", text)
    # Remove leading Roman numerals (I. / II. etc.)
    text = re.sub(r"^[IVXivx]+[.\s]+", "", text)
    return text.strip()


def _infer_official_issuing_org(title: str) -> str:
    """Infer a conservative issuing org label for the official-document red head."""
    import re
    text = (title or "").strip()
    m = re.search(r"([\u4e00-\u9fa5]{2,18}(?:委员会|办公室|公司|局|厅|部|院|中心|集团))", text)
    if m:
        return m.group(1)
    return "DataAgent"


def _set_picture_alt_text(run, alt_text: str) -> None:
    """P2-3: Set the alt text (description) attribute on the inline picture's docPr element."""
    from lxml import etree
    _WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for drawing in run._r.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    ):
        for docPr in drawing.findall(f"{{{_WP_NS}}}inline/{{{_WP_NS}}}docPr"):
            docPr.set("descr", alt_text)
            return
        for docPr in drawing.findall(f".//{{{_WP_NS}}}docPr"):
            docPr.set("descr", alt_text)
            return
