"""
WordGenerator — high-fidelity markdown → .docx renderer.

Features:
  ▸ Markdown tables (| col | col |) → native Word tables with styled header row
  ▸ **bold**, *italic*, `code` inline spans (nested-safe)
  ▸ Heading levels (## / ###)
  ▸ Bullet (-/*)  and numbered (1.) lists
  ▸ PIL Image objects, image bytes, or file paths → embedded Word pictures
  ▸ Cross-platform CJK-safe fonts: 宋体 body / 黑体 headings
  ▸ No auto-generated TOC, executive summary, or review notes

XML correctness notes:
  • To set East-Asian font on a Run you must write into w:r > w:rPr > w:rFonts,
    NOT into w:r > w:rFonts.  The helper _rpr_of() normalises this difference
    between Run objects (where Font._element is CT_R) and Style objects
    (where Font._element is CT_RPr).
  • All four rFonts attributes (ascii, hAnsi, eastAsia, cs) are set together
    so Word's style-inheritance chain doesn't accidentally pull in a fallback.
"""
from __future__ import annotations

import io
import os
import re
import tempfile
from typing import Any, Dict, List, Optional, Union


# ────────────────────────────────────────────────────────────────────────────
# Colour palette
# ────────────────────────────────────────────────────────────────────────────
_C_HEADING1 = (0x1A, 0x1A, 0x2E)
_C_HEADING2 = (0x16, 0x21, 0x3E)
_C_HEADING3 = (0x0F, 0x35, 0x60)
_C_TH_BG    = "2E4057"    # table header bg (dark blue-grey)
_C_TH_FG    = (0xFF, 0xFF, 0xFF)
_C_TITLE    = (0x1A, 0x1A, 0x2E)
_C_CAPTION  = (0x88, 0x88, 0x88)


class WordGenerator:
    """Convert structured section data (with Markdown bodies) to .docx bytes."""

    BODY_FONT    = "宋体"    # SimSun  – built into Windows / macOS / most Linux
    HEADING_FONT = "黑体"    # SimHei  – same
    BODY_PT      = 11
    LINE_PT      = 16        # line spacing (exact, in pt × 12700 twips)

    # ── public API ────────────────────────────────────────────────────────────

    async def generate(
        self,
        content: Dict[str, Any],
        template_path: Optional[str] = None,
    ) -> bytes:
        """
        content schema
        ──────────────
        {
          "title":    str,
          "subtitle": str   (optional),
          "author":   str   (optional),
          "date":     str   (optional),
          "sections": [
            {
              "heading":       str,
              "level":         int,         # 1–3
              "content":       str,         # markdown text (tables, bold, etc.)
              "bullet_points": [str],       # optional pre-parsed bullets
              "numbered_list": [str],       # optional
              "table": {                    # optional pre-parsed table
                "headers": [str],
                "rows":    [[str]]
              },
              "image_paths":   [str | bytes | PIL.Image],  # optional
            }
          ]
        }
        abstract / appendix keys are intentionally ignored.
        """
        try:
            from docx import Document
            from docx.shared import Cm, Pt

            doc = Document(
                template_path
                if (template_path and os.path.exists(template_path))
                else None,
            )
            self._setup_document(doc)
            self._add_title_block(doc, content)
            for sec in content.get("sections") or []:
                self._add_section(doc, sec)
            self._add_page_numbers(doc)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()
        except ImportError:
            return self._text_fallback(content)
        except Exception as exc:
            raise RuntimeError(f"Word文档生成失败: {exc}") from exc

    # ── document-level setup ─────────────────────────────────────────────────

    def _setup_document(self, doc) -> None:
        from docx.shared import Cm, Pt

        for sec in doc.sections:
            sec.top_margin    = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin   = Cm(3.17)
            sec.right_margin  = Cm(3.17)

        normal = doc.styles["Normal"]
        normal.font.name = self.BODY_FONT
        normal.font.size = Pt(self.BODY_PT)
        normal.paragraph_format.space_after = Pt(4)
        self._set_fonts_on_element(normal.font._element, self.BODY_FONT)

        for level, (size, rgb) in enumerate(
            [(18, _C_HEADING1), (14, _C_HEADING2), (12, _C_HEADING3)], start=1
        ):
            try:
                from docx.shared import RGBColor
                st = doc.styles[f"Heading {level}"]
                st.font.name  = self.HEADING_FONT
                st.font.size  = Pt(size)
                st.font.bold  = True
                st.font.color.rgb = RGBColor(*rgb)
                st.paragraph_format.space_before = Pt(12)
                st.paragraph_format.space_after  = Pt(4)
                self._set_fonts_on_element(st.font._element, self.HEADING_FONT)
            except Exception:
                pass

    # ── title block ──────────────────────────────────────────────────────────

    def _add_title_block(self, doc, content: Dict[str, Any]) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc.add_paragraph()   # top margin

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(content.get("title") or "报告")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(*_C_TITLE)
        self._set_fonts_on_run(run, self.HEADING_FONT)

        if content.get("subtitle"):
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(content["subtitle"])
            sr.font.size = Pt(13)
            self._set_fonts_on_run(sr, self.BODY_FONT)

        parts: List[str] = []
        if content.get("author"):
            parts.append(str(content["author"]))
        if content.get("date"):
            parts.append(str(content["date"]))
        if parts:
            mp = doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mr = mp.add_run("  ·  ".join(parts))
            mr.font.size = Pt(10)
            mr.font.color.rgb = RGBColor(*_C_CAPTION)
            self._set_fonts_on_run(mr, self.BODY_FONT)

        doc.add_page_break()

    # ── section ──────────────────────────────────────────────────────────────

    def _add_section(self, doc, section: Dict[str, Any]) -> None:
        level   = min(int(section.get("level") or 1), 3)
        heading = (section.get("heading") or "").strip()
        if heading:
            h = doc.add_heading(heading, level=level)
            for run in h.runs:
                self._set_fonts_on_run(run, self.HEADING_FONT)

        body = (section.get("content") or "").strip()
        if body:
            # Strip internal evidence citation IDs before rendering
            body = re.sub(r"\[E\d+[-\w]*\]", "", body)
            body = re.sub(r" {2,}", " ", body).strip()
            self._render_markdown(doc, body)

        for item in section.get("bullet_points") or []:
            p = doc.add_paragraph(style="List Bullet")
            self._add_inline(p, str(item))

        for item in section.get("numbered_list") or []:
            p = doc.add_paragraph(style="List Number")
            self._add_inline(p, str(item))

        # Pre-parsed table dict
        tdata = section.get("table")
        if isinstance(tdata, dict):
            hdrs = tdata.get("headers") or []
            rows = tdata.get("rows") or []
            if hdrs or rows:
                all_rows = ([hdrs] if hdrs else []) + list(rows)
                self._render_table(doc, all_rows, has_header=bool(hdrs))

        # Images: file path, bytes, or PIL Image
        for img in section.get("image_paths") or []:
            self._embed_image(doc, img)

    # ── markdown renderer ─────────────────────────────────────────────────────

    def _render_markdown(self, doc, text: str) -> None:
        """Full markdown-to-docx renderer: tables, headings, lists, bold."""
        from docx.shared import Pt

        lines = text.split("\n")
        i = 0
        while i < len(lines):
            raw = lines[i]
            line = raw.rstrip()

            # ── Markdown pipe table ──────────────────────────────────────
            if re.match(r"^\s*\|", line):
                block: List[str] = []
                while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                    block.append(lines[i])
                    i += 1
                self._render_md_table(doc, block)
                continue

            # ── ATX headings ─────────────────────────────────────────────
            m = re.match(r"^(#{1,3})\s+(.*)", line)
            if m:
                depth = len(m.group(1))
                h = doc.add_heading(m.group(2).strip(), level=min(depth, 3))
                for run in h.runs:
                    self._set_fonts_on_run(run, self.HEADING_FONT)
                i += 1
                continue

            # ── Bullet list ──────────────────────────────────────────────
            m = re.match(r"^(\s*)[-*]\s+(.*)", line)
            if m:
                p = doc.add_paragraph(style="List Bullet")
                self._add_inline(p, m.group(2).strip())
                i += 1
                continue

            # ── Numbered list ────────────────────────────────────────────
            m = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
            if m:
                p = doc.add_paragraph(style="List Number")
                self._add_inline(p, m.group(2).strip())
                i += 1
                continue

            # ── Blank line → small spacer ────────────────────────────────
            if not line.strip():
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # ── Normal paragraph ─────────────────────────────────────────
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            self._add_inline(p, line.strip())
            i += 1

    # ── markdown table ────────────────────────────────────────────────────────

    def _render_md_table(self, doc, lines: List[str]) -> None:
        """Parse GitHub-flavor markdown table lines → Word table."""
        rows: List[List[str]] = []
        for line in lines:
            stripped = line.strip().strip("|")
            if re.match(r"^[\s\-:|]+$", stripped):
                # Separator row — marks end of header
                continue
            cells = [c.strip() for c in stripped.split("|")]
            rows.append(cells)
        if not rows:
            return
        self._render_table(doc, rows, has_header=True)

    def _render_table(
        self, doc, rows: List[List[str]], *, has_header: bool = True
    ) -> None:
        from docx.shared import RGBColor, Pt
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        n_rows = len(rows)
        if n_cols == 0:
            return

        try:
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Table Grid"

            for ri, row in enumerate(rows):
                is_header = has_header and ri == 0
                for ci in range(n_cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    cell = table.cell(ri, ci)
                    # Clear any default text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ""
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    self._add_inline(p, cell_text, bold_all=is_header)
                    if is_header:
                        self._set_cell_bg(cell, _C_TH_BG)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(*_C_TH_FG)

            doc.add_paragraph()   # breathing room after table
        except Exception:
            pass

    # ── inline text ──────────────────────────────────────────────────────────

    _INLINE_PATTERN = re.compile(
        r"(\*\*(?P<bold>[^*]+)\*\*"
        r"|\*(?P<italic>[^*]+)\*"
        r"|`(?P<code>[^`]+)`)"
    )

    def _add_inline(
        self, para, text: str, *, bold_all: bool = False
    ) -> None:
        """Render **bold**, *italic*, `code` inline spans into a paragraph."""
        pos = 0
        for m in self._INLINE_PATTERN.finditer(text):
            # plain text before match
            if m.start() > pos:
                run = para.add_run(text[pos : m.start()])
                if bold_all:
                    run.bold = True
                self._set_fonts_on_run(run, self.BODY_FONT)

            if m.group("bold"):
                run = para.add_run(m.group("bold"))
                run.bold = True
                self._set_fonts_on_run(run, self.BODY_FONT)
            elif m.group("italic"):
                run = para.add_run(m.group("italic"))
                run.italic = True
                self._set_fonts_on_run(run, self.BODY_FONT)
            elif m.group("code"):
                run = para.add_run(m.group("code"))
                run.font.name = "Courier New"
                try:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                except Exception:
                    pass
            pos = m.end()

        # trailing plain text
        if pos < len(text):
            run = para.add_run(text[pos:])
            if bold_all:
                run.bold = True
            self._set_fonts_on_run(run, self.BODY_FONT)

    # ── image embedding ───────────────────────────────────────────────────────

    def _embed_image(self, doc, img: Any) -> None:
        """Embed a PIL.Image, raw bytes, or file path as a centred picture."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile

        tmp_path: Optional[str] = None
        try:
            # PIL Image object
            if hasattr(img, "save"):
                tmp = tempfile.NamedTemporaryFile(
                    suffix=".png", delete=False
                )
                img.save(tmp, format="PNG")
                tmp.close()
                tmp_path = tmp.name
                img_path = tmp_path
            # Raw bytes
            elif isinstance(img, (bytes, bytearray)):
                tmp = tempfile.NamedTemporaryFile(
                    suffix=".png", delete=False
                )
                tmp.write(img)
                tmp.close()
                tmp_path = tmp.name
                img_path = tmp_path
            # File path
            elif isinstance(img, str) and os.path.exists(img):
                img_path = img
            else:
                return

            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(img_path, width=Inches(5.8))
        except Exception:
            pass
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

    # ── XML helpers ───────────────────────────────────────────────────────────

    def _set_fonts_on_run(self, run, font_name: str) -> None:
        """
        Correctly set all four w:rFonts attributes on a Run's rPr.

        Run.font._element  is  w:r  (CT_R) — we must descend to rPr first.
        Style.font._element is w:rPr (CT_RPr) — rFonts is a direct child.
        """
        try:
            from docx.oxml.ns import qn
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), font_name)
        except Exception:
            # Graceful fallback: at least set the high-level name
            try:
                run.font.name = font_name
            except Exception:
                pass

    def _set_fonts_on_element(self, rpr_element, font_name: str) -> None:
        """
        Set font names on a CT_RPr element directly (for style-level fonts).
        Style.font._element is already w:rPr — no descent needed.
        """
        try:
            from docx.oxml.ns import qn
            rFonts = rpr_element.get_or_add_rFonts()
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), font_name)
        except Exception:
            pass

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Remove any existing shd
            for old in tcPr.findall(qn("w:shd")):
                tcPr.remove(old)
            tcPr.append(shd)
        except Exception:
            pass

    def _add_page_numbers(self, doc) -> None:
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            for sec in doc.sections:
                footer = sec.footer
                fp = (
                    footer.paragraphs[0]
                    if footer.paragraphs
                    else footer.add_paragraph()
                )
                fp.clear()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = fp.add_run()
                for ftype in ("begin", None, "end"):
                    if ftype is None:
                        instr = OxmlElement("w:instrText")
                        instr.text = " PAGE "
                        run._r.append(instr)
                    else:
                        fc = OxmlElement("w:fldChar")
                        fc.set(qn("w:fldCharType"), ftype)
                        run._r.append(fc)
        except Exception:
            pass

    # ── fallback (no python-docx) ─────────────────────────────────────────────

    def _text_fallback(self, content: Dict[str, Any]) -> bytes:
        lines = [f"# {content.get('title', 'Document')}\n"]
        for sec in content.get("sections") or []:
            h = "#" * sec.get("level", 2)
            lines.append(f"\n{h} {sec.get('heading', '')}\n")
            if sec.get("content"):
                lines.append(sec["content"] + "\n")
        return "\n".join(lines).encode("utf-8")
