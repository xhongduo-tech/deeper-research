"""DocumentGraph — Layout-aware document structure extraction.

SOTA gap closed: baseline `file_parser.py` produces flat text streams that lose
document topology — there's no way to know "this paragraph belongs to Section
3.2", or "this is a real table with 5 columns" vs "this is body text". That
makes section-targeted retrieval, citation-bearing extraction, and table-as-table
analysis impossible.

This module reads PDF/DOCX/XLSX and emits a typed element graph:

    Document
    ├── Heading (level=1, "1. 总体概览")
    │   ├── Paragraph ("…")
    │   ├── Table (rows=15, cols=4, df=...)
    │   └── Heading (level=2, "1.1 业务规模")
    │       └── Paragraph ("…")
    ├── Heading (level=1, "2. 风险分析")
    │   └── List (items=["…", "…"])
    └── Figure (caption="图1 各分行贷款余额")

Why this matters:
  • Section-aware RAG chunking (chunk by H1/H2 boundaries, not arbitrary chars)
  • Table-as-table extraction (preserve relations for downstream SQL/join)
  • Outline reconstruction ("section 3.2 says X")
  • Layout-aware citation ("…[文档.pdf §2.1]")

Backend uses Docling-inspired primitives but with zero external dependencies
beyond what's already in requirements.txt (python-docx, pdfplumber, openpyxl).
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal, Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Element types — the building blocks of a document graph
# ─────────────────────────────────────────────────────────────────────────────

ElementKind = Literal[
    "heading",      # section title with .level (1=H1, 2=H2, …)
    "paragraph",    # body prose
    "list_item",    # bullet or numbered list item
    "table",        # structured table with .df (pandas DataFrame)
    "figure",       # embedded image / chart with optional .caption
    "code",         # code block
    "footnote",     # footnote / endnote
    "caption",      # standalone caption
    "footer",       # repeated page footer (filtered from main flow)
    "header",       # repeated page header (filtered from main flow)
]


@dataclass
class DocumentElement:
    """One typed unit in the document graph."""
    kind: ElementKind
    text: str = ""
    level: int = 0                      # for headings (1..6) or list nesting
    page: int = 0                       # 1-indexed page number
    bbox: tuple | None = None           # (x0, y0, x1, y1) when available
    table_data: Optional["TableData"] = None
    metadata: dict = field(default_factory=dict)
    parent_index: int = -1              # index of parent heading in elements[]

    @property
    def is_structural(self) -> bool:
        return self.kind in ("heading", "table", "figure", "list_item")


@dataclass
class TableData:
    """Structured table representation — relational, not stringified."""
    rows: list[list[str]] = field(default_factory=list)
    header: list[str] = field(default_factory=list)
    caption: str = ""
    page: int = 0

    @property
    def n_rows(self) -> int:
        return len(self.rows)

    @property
    def n_cols(self) -> int:
        return len(self.header) or (len(self.rows[0]) if self.rows else 0)

    def to_dataframe(self):
        import pandas as pd
        if not self.rows:
            return pd.DataFrame(columns=self.header)
        if self.header and all(len(r) == len(self.header) for r in self.rows[:5]):
            return pd.DataFrame(self.rows, columns=self.header)
        return pd.DataFrame(self.rows)

    def to_markdown(self) -> str:
        if not self.rows and not self.header:
            return ""
        out = []
        if self.caption:
            out.append(f"**{self.caption}**\n")
        hdr = self.header if self.header else [f"col_{i+1}" for i in range(self.n_cols)]
        out.append("| " + " | ".join(hdr) + " |")
        out.append("| " + " | ".join("---" for _ in hdr) + " |")
        for row in self.rows:
            cells = [str(c).replace("|", "｜") for c in row]
            out.append("| " + " | ".join(cells) + " |")
        return "\n".join(out)


@dataclass
class DocumentGraph:
    """The whole document — flat element list + helpers to navigate."""
    elements: list[DocumentElement] = field(default_factory=list)
    source_path: str = ""
    source_type: str = ""               # pdf | docx | xlsx
    page_count: int = 0
    title: str = ""

    # ── Navigation helpers ────────────────────────────────────────────────

    def headings(self) -> list[DocumentElement]:
        return [e for e in self.elements if e.kind == "heading"]

    def tables(self) -> list[DocumentElement]:
        return [e for e in self.elements if e.kind == "table"]

    def figures(self) -> list[DocumentElement]:
        return [e for e in self.elements if e.kind == "figure"]

    def outline(self, max_level: int = 3) -> list[dict]:
        """Return heading-only outline as [{level, text, index}, ...]."""
        return [
            {"level": e.level, "text": e.text, "index": i}
            for i, e in enumerate(self.elements)
            if e.kind == "heading" and e.level <= max_level
        ]

    def section_text(self, heading_index: int) -> str:
        """All non-heading text under a given heading until the next heading
        of equal or higher level."""
        if heading_index < 0 or heading_index >= len(self.elements):
            return ""
        anchor = self.elements[heading_index]
        if anchor.kind != "heading":
            return ""
        chunks: list[str] = []
        for e in self.elements[heading_index + 1:]:
            if e.kind == "heading" and e.level <= anchor.level:
                break
            if e.kind == "paragraph":
                chunks.append(e.text)
            elif e.kind == "list_item":
                chunks.append(f"• {e.text}")
            elif e.kind == "table" and e.table_data:
                chunks.append(e.table_data.to_markdown())
        return "\n\n".join(chunks)

    def to_sectioned_text(self) -> str:
        """Render the whole document as section-prefixed text for RAG chunking.

        Each chunk is preceded by its breadcrumb path
        (e.g. "[1 概述 > 1.1 背景]") so retrieval can attribute findings."""
        lines: list[str] = []
        path: list[str] = []
        for e in self.elements:
            if e.kind == "heading":
                # Adjust breadcrumb to current level
                while len(path) >= e.level:
                    path.pop()
                path.append(e.text)
                lines.append(f"\n{'#' * e.level} {e.text}")
            elif e.kind == "paragraph":
                lines.append(e.text)
            elif e.kind == "list_item":
                indent = "  " * max(0, e.level - 1)
                lines.append(f"{indent}• {e.text}")
            elif e.kind == "table" and e.table_data:
                lines.append("\n" + e.table_data.to_markdown())
            elif e.kind == "figure" and e.text:
                lines.append(f"[图: {e.text}]")
        return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX extraction — uses python-docx Styles for true heading detection
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx_graph(path: str) -> DocumentGraph:
    """Build a DocumentGraph from a DOCX file using Style names."""
    graph = DocumentGraph(source_path=path, source_type="docx")
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; falling back to plain text")
        return graph

    try:
        doc = Document(path)
    except Exception as exc:
        logger.warning("DOCX parse failed: %s", exc)
        return graph

    # Title (from core properties)
    try:
        graph.title = doc.core_properties.title or ""
    except Exception:
        pass

    last_heading_idx_per_level: dict[int, int] = {}

    # Iterate body elements in document order (paragraphs + tables interleaved)
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    # python-docx doesn't natively interleave — walk XML children for order
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1] if "}" in child.tag else child.tag
        if tag == "p":
            try:
                para = next(para_iter)
            except StopIteration:
                continue
            text = (para.text or "").strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""

            # Heading detection: Style "Heading 1" / "标题 1"
            heading_match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style_name)
            if heading_match:
                level = int(heading_match.group(1) or heading_match.group(2))
                parent_idx = -1
                for lvl in range(level - 1, 0, -1):
                    if lvl in last_heading_idx_per_level:
                        parent_idx = last_heading_idx_per_level[lvl]
                        break
                el = DocumentElement(
                    kind="heading",
                    text=text,
                    level=level,
                    parent_index=parent_idx,
                )
                graph.elements.append(el)
                last_heading_idx_per_level[level] = len(graph.elements) - 1
                # Clear deeper levels
                for lvl in list(last_heading_idx_per_level.keys()):
                    if lvl > level:
                        del last_heading_idx_per_level[lvl]
                continue

            # List item detection
            if "list" in style_name or style_name.startswith("bullet"):
                level = 1  # python-docx doesn't expose list-level easily
                graph.elements.append(DocumentElement(
                    kind="list_item",
                    text=text,
                    level=level,
                    parent_index=last_heading_idx_per_level.get(
                        max(last_heading_idx_per_level.keys(), default=0), -1
                    ),
                ))
                continue

            # Regular paragraph
            parent_idx = last_heading_idx_per_level.get(
                max(last_heading_idx_per_level.keys(), default=0), -1
            )
            graph.elements.append(DocumentElement(
                kind="paragraph",
                text=text,
                parent_index=parent_idx,
            ))

        elif tag == "tbl":
            try:
                table = next(table_iter)
            except StopIteration:
                continue
            try:
                rows: list[list[str]] = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if any(c for c in cells):
                        rows.append(cells)
                if not rows:
                    continue
                # Heuristic: first row is header if it has no numbers
                first_row_is_header = all(
                    not re.search(r"\d", c) for c in rows[0]
                ) and len(rows) > 1
                header = rows[0] if first_row_is_header else []
                body_rows = rows[1:] if first_row_is_header else rows
                graph.elements.append(DocumentElement(
                    kind="table",
                    text=f"Table ({len(rows)} rows × {len(rows[0])} cols)",
                    table_data=TableData(rows=body_rows, header=header),
                    parent_index=last_heading_idx_per_level.get(
                        max(last_heading_idx_per_level.keys(), default=0), -1
                    ),
                ))
            except Exception as exc:
                logger.debug("Table parse error: %s", exc)

    return graph


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — uses pdfplumber bounding boxes + font-size heading detection
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf_graph(path: str) -> DocumentGraph:
    """Build a DocumentGraph from a PDF using pdfplumber.

    Heading detection: font size-based — a text run with significantly larger
    font than the body baseline is treated as a heading. Level is derived from
    the rank of distinct heading sizes (largest = level 1).
    """
    graph = DocumentGraph(source_path=path, source_type="pdf")
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed; document graph empty")
        return graph

    try:
        with pdfplumber.open(path) as pdf:
            try:
                graph.title = (pdf.metadata or {}).get("Title", "") or ""
            except Exception:
                pass
            graph.page_count = len(pdf.pages)

            # ── Pass 1: collect all text spans with font sizes ──
            all_spans: list[dict] = []
            for page_idx, page in enumerate(pdf.pages, start=1):
                try:
                    # Get character-level data for font analysis
                    chars = page.chars or []
                    if not chars:
                        continue
                    # Group chars into lines by y-coordinate
                    lines = _group_chars_into_lines(chars)
                    for line in lines:
                        if not line["text"].strip():
                            continue
                        all_spans.append({
                            "text": line["text"].strip(),
                            "size": line["avg_size"],
                            "page": page_idx,
                            "bbox": line["bbox"],
                            "is_bold": line.get("is_bold", False),
                        })
                except Exception as exc:
                    logger.debug("Page %d char extraction failed: %s", page_idx, exc)

            # ── Pass 2: determine body font size baseline ──
            if not all_spans:
                # Empty PDF — fall back to plain text
                return graph

            # Body size = mode of sizes (most common)
            from collections import Counter
            size_counter = Counter(round(s["size"], 1) for s in all_spans)
            body_size = size_counter.most_common(1)[0][0] if size_counter else 10.0

            # Identify heading sizes (>= 1.2x body) and rank them
            heading_sizes = sorted({
                round(s["size"], 1) for s in all_spans
                if round(s["size"], 1) >= body_size * 1.2
            }, reverse=True)
            heading_level_map = {sz: lvl + 1 for lvl, sz in enumerate(heading_sizes[:4])}

            # ── Pass 3: emit elements ──
            last_heading_idx_per_level: dict[int, int] = {}
            for span in all_spans:
                size_round = round(span["size"], 1)
                if size_round in heading_level_map:
                    level = heading_level_map[size_round]
                    parent_idx = -1
                    for lvl in range(level - 1, 0, -1):
                        if lvl in last_heading_idx_per_level:
                            parent_idx = last_heading_idx_per_level[lvl]
                            break
                    graph.elements.append(DocumentElement(
                        kind="heading",
                        text=span["text"],
                        level=level,
                        page=span["page"],
                        bbox=span["bbox"],
                        parent_index=parent_idx,
                    ))
                    last_heading_idx_per_level[level] = len(graph.elements) - 1
                    for lvl in list(last_heading_idx_per_level.keys()):
                        if lvl > level:
                            del last_heading_idx_per_level[lvl]
                else:
                    parent_idx = last_heading_idx_per_level.get(
                        max(last_heading_idx_per_level.keys(), default=0), -1
                    )
                    graph.elements.append(DocumentElement(
                        kind="paragraph",
                        text=span["text"],
                        page=span["page"],
                        bbox=span["bbox"],
                        parent_index=parent_idx,
                    ))

            # ── Pass 4: extract tables via pdfplumber's find_tables ──
            for page_idx, page in enumerate(pdf.pages, start=1):
                try:
                    tables = page.find_tables() or []
                    for tbl in tables:
                        rows = tbl.extract()
                        if not rows or len(rows) < 2:
                            continue
                        # Clean cells
                        rows = [
                            [(c or "").strip().replace("\n", " ") for c in row]
                            for row in rows
                            if any((c or "").strip() for c in row)
                        ]
                        if not rows:
                            continue
                        first_row_is_header = all(
                            not re.search(r"\d", c) for c in rows[0]
                        ) and len(rows) > 1
                        header = rows[0] if first_row_is_header else []
                        body_rows = rows[1:] if first_row_is_header else rows
                        graph.elements.append(DocumentElement(
                            kind="table",
                            text=f"Table p{page_idx} ({len(rows)} × {len(rows[0])})",
                            page=page_idx,
                            bbox=tbl.bbox,
                            table_data=TableData(
                                rows=body_rows,
                                header=header,
                                page=page_idx,
                            ),
                        ))
                except Exception as exc:
                    logger.debug("Table extraction on page %d failed: %s", page_idx, exc)
    except Exception as exc:
        logger.warning("PDF graph extraction failed for %s: %s", path, exc)

    return graph


def _group_chars_into_lines(chars: list[dict]) -> list[dict]:
    """Group pdfplumber char dicts into line dicts by y-coordinate.

    Returns: [{"text": str, "avg_size": float, "bbox": (x0,y0,x1,y1), "is_bold": bool}]
    """
    if not chars:
        return []
    lines: list[dict] = []
    current_y = None
    current_chars: list[dict] = []

    # Sort by y desc then x asc (top-to-bottom, left-to-right)
    sorted_chars = sorted(chars, key=lambda c: (-round(c.get("top", 0), 0), c.get("x0", 0)))

    for ch in sorted_chars:
        y = round(ch.get("top", 0), 0)
        if current_y is None or abs(y - current_y) <= 2:
            current_chars.append(ch)
            current_y = y if current_y is None else current_y
        else:
            lines.append(_finalize_line(current_chars))
            current_chars = [ch]
            current_y = y
    if current_chars:
        lines.append(_finalize_line(current_chars))
    return [l for l in lines if l["text"]]


def _finalize_line(chars: list[dict]) -> dict:
    if not chars:
        return {"text": "", "avg_size": 0, "bbox": (0, 0, 0, 0), "is_bold": False}
    text = "".join(c.get("text", "") for c in chars)
    sizes = [c.get("size", 0) for c in chars if c.get("size")]
    avg_size = sum(sizes) / len(sizes) if sizes else 0
    x0 = min(c.get("x0", 0) for c in chars)
    x1 = max(c.get("x1", 0) for c in chars)
    y0 = min(c.get("top", 0) for c in chars)
    y1 = max(c.get("bottom", 0) for c in chars)
    is_bold = any("Bold" in str(c.get("fontname", "")) for c in chars)
    return {
        "text": text,
        "avg_size": avg_size,
        "bbox": (x0, y0, x1, y1),
        "is_bold": is_bold,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def build_graph(path: str | Path) -> DocumentGraph:
    """Build a DocumentGraph from any supported file type.

    Currently supports PDF and DOCX. Other types return empty graph (callers
    should fall back to the flat-text extractor in `file_parser.py`).
    """
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in (".pdf",):
        return extract_pdf_graph(str(p))
    if suffix in (".docx", ".doc"):
        return extract_docx_graph(str(p))
    return DocumentGraph(source_path=str(p), source_type=suffix.lstrip("."))


def summarize_graph(graph: DocumentGraph) -> dict:
    """Compact JSON-safe summary for telemetry / WS broadcasting."""
    return {
        "source_type": graph.source_type,
        "page_count": graph.page_count,
        "title": graph.title,
        "n_elements": len(graph.elements),
        "n_headings": len(graph.headings()),
        "n_tables": len(graph.tables()),
        "n_figures": len(graph.figures()),
        "outline_depth": max((e.level for e in graph.headings()), default=0),
        "outline": graph.outline(max_level=2)[:12],
    }
